import os, sys, traceback, statistics
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))

#小数点处理部分以及N/A处理
#四舍五入到两位
def fmt2(v):
    if v is None: return ''
    return f'{v:.2f}' if isinstance(v, (int, float)) else str(v)
#取整，比如行业排名
def fmt0(v):
    if v is None: return ''
    return str(round(v)) if isinstance(v, (int, float)) else str(v)
#百分比
def fmt_pct(v):
    if v is None: return ''
    return f'{v*100:.2f}%' if isinstance(v, (int, float)) else str(v)
#字符串，比如股票编号
def fmt_str(v):
    return '' if v is None else str(v)
#日期
def fmt_date(v):
    return v.strftime('%Y-%m-%d') if isinstance(v, datetime) else (str(v) if v else '')

def xl(ws, r, c):
    try: return ws.cell(r, c).value
    except: return None

# Word tools
#单元格里写内容
def set_cell(cell, text, font_size=9, bold=False, align='center'):
    if cell is None: return
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text is not None else '')
    run.font.name = '宋体'; run.font.size = Pt(font_size); run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

#给表格加边框
def add_borders(table):
    xml = (f'<w:tblBorders {nsdecls("w")}>'
        + ''.join(f'<w:{s} w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                  for s in ['top','left','bottom','right','insideH','insideV'])
        + '</w:tblBorders>')
    pr = table._tbl.tblPr
    if pr is None: pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); table._tbl.insert(0, pr)
    for old in pr.findall(qn('w:tblBorders')): pr.remove(old)
    pr.append(parse_xml(xml))

#给单元格加底色
def shade_cell(cell, color='D9D9D9'):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None: tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>'); tc.insert(0, tcPr)
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

#指定位置创建表格（先创再移）
def make_table(doc, rows, cols, after_el, col_widths=None, compact=False):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = False
    add_borders(tbl)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells): row.cells[i].width = Cm(w)
    if compact:
        m = parse_xml(f'<w:tblCellMar {nsdecls("w")}><w:top w:w="0" w:type="dxa"/>'
            '<w:left w:w="30" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/>'
            '<w:right w:w="30" w:type="dxa"/></w:tblCellMar>')
        for old in tbl._tbl.tblPr.findall(qn('w:tblCellMar')): tbl._tbl.tblPr.remove(old)
        tbl._tbl.tblPr.append(m)
    doc.element.body.remove(tbl._element)
    after_el.addnext(tbl._element)
    return tbl

#特定位置插段落
def make_para(doc, text, after_el, font_size=9, bold=False):
    p = doc.add_paragraph(); run = p.add_run(text)
    run.font.name = '宋体'; run.font.size = Pt(font_size); run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.element.body.remove(p._element); after_el.addnext(p._element)
    return p

def find_para(doc, kw):
    for p in doc.paragraphs:
        if kw in p.text: return p._element
    return None

#清空内容，有空间才能插入表格
def clear_between(doc, start_kw, end_kw):
    body = doc.element.body
    s = find_para(doc, start_kw); e = find_para(doc, end_kw) if end_kw else None
    if s is None: return None
    rm, inside = [], False
    for ch in body.iterchildren():
        if ch is s: inside = True; continue
        if e is not None and ch is e: break
        if inside: rm.append(ch)
    for el in rm: body.remove(el)
    return s

#找到文件夹里第一个word当作模板，和第一个excel当作数据库
def get_files():
    ex = wd = None
    for f in sorted(os.listdir(BASE_DIR)):
        if f.startswith('~$') or '合并完成' in f: continue
        if f.endswith('.xlsx') and not ex: ex = os.path.join(BASE_DIR, f)
        elif f.endswith('.docx') and not wd: wd = os.path.join(BASE_DIR, f)
    return ex, wd


#本次交易要素概要表 → "二、申请项目要素"下方，去掉第2笔/第3笔/合计三列
#原大表r0-22替换成新表，r23-35（其他项目情况）保留为独立表格
def fill_deal_summary(doc, wb_val, wb_fmt):
    if '本次要素' not in wb_val.sheetnames: return print('没有"本次要素"sheet')
    ws = wb_val['本次要素']
    wf = wb_fmt['本次要素']

    #找"二、申请项目要素"标题
    anchor = find_para(doc, '申请项目要素')
    if anchor is None: return print('未找到"二、申请项目要素"段落')

    #往后扫到第一个表，拆分保留其他项目情况(r23+)
    import copy
    el = anchor.getnext()
    old_tbl_el = None
    while el is not None:
        if el.tag.split('}')[-1] == 'tbl': old_tbl_el = el; break
        el = el.getnext()
    if old_tbl_el is None: return print('未找到申请项目要素表格')

    #深拷贝原表，删前23行只留其他项目情况
    keep_tbl = copy.deepcopy(old_tbl_el)
    all_rows = keep_tbl.findall(qn('w:tr'))
    for tr in all_rows[:23]: keep_tbl.remove(tr)  # 保留r23-r35

    #删原表
    doc.element.body.remove(old_tbl_el)

    #根据Excel number_format自动格式化值
    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return ''
        if isinstance(v, datetime): return v.strftime('%Y/%m/%d')
        if isinstance(v, str): return v
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            dec = 2 if '0.00%' in nf else 0
            return f'{v*100:.{dec}f}%'
        if '#,##0' in nf:
            return f'{v:,.2f}' if '0.00' in nf else f'{v:,.0f}'
        if '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    #行号：r15-32主体，r33子标题，r34-37交易后
    MAIN = list(range(15, 33))
    POST = list(range(34, 38))
    nrows = 1 + len(MAIN) + 1 + len(POST)

    tbl = make_table(doc, nrows, 4, anchor, col_widths=[3.5, 3.0, 3.5, 3.5])

    #表头：灰底，基本信息合并两格
    for ci in range(4): shade_cell(tbl.cell(0, ci))
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    set_cell(tbl.cell(0, 0), '基本信息', bold=True)
    set_cell(tbl.cell(0, 2), '交易要素', bold=True)
    set_cell(tbl.cell(0, 3), '第1笔', bold=True)

    #主体数据 r15-32
    for i, r in enumerate(MAIN):
        ri = 1 + i
        set_cell(tbl.cell(ri, 0), fv(r, 2), align='left')
        set_cell(tbl.cell(ri, 1), fv(r, 3))
        set_cell(tbl.cell(ri, 2), fv(r, 4), align='left')
        set_cell(tbl.cell(ri, 3), fv(r, 5))

    #子标题行：交易完成后项目融资情况，灰底合并
    sub_ri = 1 + len(MAIN)
    for ci in range(4): shade_cell(tbl.cell(sub_ri, ci))
    tbl.cell(sub_ri, 0).merge(tbl.cell(sub_ri, 3))
    set_cell(tbl.cell(sub_ri, 0), '交易完成后项目融资情况', bold=True, align='left')

    #交易后数据 r34-37
    for i, r in enumerate(POST):
        ri = sub_ri + 1 + i
        set_cell(tbl.cell(ri, 0), fv(r, 2), align='left')
        set_cell(tbl.cell(ri, 1), fv(r, 3))
        set_cell(tbl.cell(ri, 2), fv(r, 4), align='left')
        set_cell(tbl.cell(ri, 3), fv(r, 5))

    #把保留的其他项目情况表格插到新表后面
    tbl._element.addnext(keep_tbl)

    print('本次要素 (申请项目要素)')


#多行内容写入同一个单元格（关联股东情况等需要换行的场景）
def set_cell_lines(cell, lines, font_size=9, bold=False, align='left'):
    if cell is None: return
    cell.text = ''
    al = WD_ALIGN_PARAGRAPH.LEFT if align == 'left' else WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(lines):
        if not line: continue
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = al
        run = p.add_run(str(line))
        run.font.name = '宋体'; run.font.size = Pt(font_size); run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


#持股情况表 1、持股情况 下方，删旧表建新4列表
def fill_shareholding(doc, wb_val, wb_fmt):
    if '本次要素' not in wb_val.sheetnames: return print('没有"本次要素"sheet')
    ws = wb_val['本次要素']
    wf = wb_fmt['本次要素']

    #找 1、持股情况段落，删紧跟的旧表
    anchor = find_para(doc, '1、持股情况')
    if anchor is None: return print('未找到"1、持股情况"段落')
    el = anchor.getnext()
    while el is not None:
        if el.tag.split('}')[-1] == 'tbl': doc.element.body.remove(el); break
        el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return ''
        if isinstance(v, datetime): return v.strftime('%Y/%m/%d')
        if isinstance(v, str): return v
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            dec = 2 if '0.00%' in nf else 0
            return f'{v*100:.{dec}f}%'
        if '#,##0' in nf:
            return f'{v:,.2f}' if '0.00' in nf else f'{v:,.0f}'
        if '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    tbl = make_table(doc, 16, 4, anchor, col_widths=[3.5, 2.5, 3.5, 4.0])

    #表头：灰底，左右各合并两格
    for ci in range(4): shade_cell(tbl.cell(0, ci))
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    set_cell(tbl.cell(0, 0), '融资方持股情况', bold=True)
    tbl.cell(0, 2).merge(tbl.cell(0, 3))
    set_cell(tbl.cell(0, 2), '股份减持', bold=True)

    #主体左侧 c2/c3 逐行填入 (r64-73 → 表行1-10)
    for i, r in enumerate(range(64, 74)):
        ri = 1 + i
        set_cell(tbl.cell(ri, 0), fv(r, 2), align='left')
        set_cell(tbl.cell(ri, 1), fv(r, 3))

    #右侧：前两行单独填 (r64-65)
    set_cell(tbl.cell(1, 2), fv(64, 4), align='left')
    set_cell(tbl.cell(1, 3), fv(64, 5))
    set_cell(tbl.cell(2, 2), fv(65, 4), align='left')
    set_cell(tbl.cell(2, 3), fv(65, 5))

    #右侧：关联股东情况，r66-69合并 (表行3-6)
    tbl.cell(3, 2).merge(tbl.cell(6, 2))
    set_cell(tbl.cell(3, 2), fv(66, 4), align='left')
    tbl.cell(3, 3).merge(tbl.cell(6, 3))
    set_cell(tbl.cell(3, 3), fv(66, 5), align='left')

    #右侧：股份减持受限描述，r70-72合并 (表行7-9)
    tbl.cell(7, 2).merge(tbl.cell(9, 2))
    set_cell(tbl.cell(7, 2), fv(70, 4), align='left')
    tbl.cell(7, 3).merge(tbl.cell(9, 3))
    desc = [fv(70, 5), fv(71, 5)]
    set_cell_lines(tbl.cell(7, 3), [d for d in desc if d])

    #右侧：质押股份处置周期 r73 (表行10)
    set_cell(tbl.cell(10, 2), fv(73, 4), align='left')
    set_cell(tbl.cell(10, 3), fv(73, 5))

    #子标题：交易完成后项目融资情况（证金公司报送）(表行11)
    for ci in range(4): shade_cell(tbl.cell(11, ci))
    tbl.cell(11, 0).merge(tbl.cell(11, 3))
    set_cell(tbl.cell(11, 0), fv(74, 2), bold=True, align='left')

    #r75-76 (表行12-13)
    for i, r in enumerate([75, 76]):
        ri = 12 + i
        set_cell(tbl.cell(ri, 0), fv(r, 2), align='left')
        set_cell(tbl.cell(ri, 1), fv(r, 3))
        set_cell(tbl.cell(ri, 2), fv(r, 4), align='left')
        set_cell(tbl.cell(ri, 3), fv(r, 5))

    #子标题：第一大股东高比例质押及受限股（风控指标报送）(表行14)
    for ci in range(4): shade_cell(tbl.cell(14, ci))
    tbl.cell(14, 0).merge(tbl.cell(14, 3))
    set_cell(tbl.cell(14, 0), fv(77, 2), bold=True, align='left')

    #r78 (表行15)
    set_cell(tbl.cell(15, 0), fv(78, 2), align='left')
    set_cell(tbl.cell(15, 1), fv(78, 3))
    set_cell(tbl.cell(15, 2), fv(78, 4), align='left')
    set_cell(tbl.cell(15, 3), fv(78, 5))

    print('持股情况 (本次要素)')


#控股股东减持受限情况表 2、破发、破净 段落后，删旧表建新2列表
def fill_restrict(doc, wb_val, wb_fmt):
    if '本次要素' not in wb_val.sheetnames: return print('没有"本次要素"sheet')
    ws = wb_val['本次要素']
    wf = wb_fmt['本次要素']

    #找 破发 段落，找到旧表并删除
    anchor = None
    for p in doc.paragraphs:
        if '破发' in p.text and '破净' in p.text:
            anchor = p._element; break
    if anchor is None: return print('未找到"破发、破净"段落')
    prev_el = anchor; el = anchor.getnext()
    while el is not None:
        if el.tag.split('}')[-1] == 'tbl':
            doc.element.body.remove(el); break
        prev_el = el; el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return ''
        if isinstance(v, datetime): return v.strftime('%Y/%m/%d')
        if isinstance(v, str): return v
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            dec = 2 if '0.00%' in nf else 0
            return f'{v*100:.{dec}f}%'
        if '#,##0' in nf:
            return f'{v:,.2f}' if '0.00' in nf else f'{v:,.0f}'
        if '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    # 9行：表头1(r81) + 数据8(r82-89)，2列
    ROWS = list(range(82, 90))
    tbl = make_table(doc, 1 + len(ROWS), 2, prev_el, col_widths=[7.0, 6.5])

    #表头：灰底合并
    for ci in range(2): shade_cell(tbl.cell(0, ci))
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    set_cell(tbl.cell(0, 0), fv(81, 2), bold=True)

    #数据行：c2=标签，c5=值
    for i, r in enumerate(ROWS):
        ri = 1 + i
        set_cell(tbl.cell(ri, 0), fv(r, 2), align='left')
        set_cell(tbl.cell(ri, 1), fv(r, 5))

    print('控股股东减持受限 (本次要素)')


#标的证券基本面，往模板已有表格里填值，含主营业务从收入构成提取产品名
def fill_basic_info(doc, wb_val, wb_fmt):
    if '标的券概要' not in wb_val.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb_val['标的券概要']
    wf = wb_fmt['标的券概要']

    #找表
    tbl = None
    for t in doc.tables:
        for row in t.rows[:1]:
            for c in row.cells:
                if '标的证券基本面' in c.text: tbl = t; break
            if tbl: break
        if tbl: break
    if tbl is None: return print('未找到"标的证券基本面"表格')

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return ''
        if isinstance(v, datetime): return v.strftime('%Y/%m/%d')
        if isinstance(v, str): return v
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            dec = 2 if '0.00%' in nf else 0
            return f'{v*100:.{dec}f}%'
        if '#,##0' in nf:
            return f'{v:,.2f}' if '0.00' in nf else f'{v:,.0f}'
        if '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    def sc(r, c):
        try: return tbl.cell(r, c) if 0 <= r < len(tbl.rows) and 0 <= c < len(tbl.columns) else None
        except: return None

    #r1：信息披露考评结果，替换年份，填等级
    year_text = fmt_str(xl(ws, 42, 2))          # "上市公司2023年信息披露考评结果"
    set_cell(sc(1, 0), year_text, bold=True, align='left')
    set_cell(sc(1, 4), fv(42, 6), bold=True)     # 等级 B

    #r2：主营业务，从收入构成（按产品）提取产品名、去掉百分比
    raw = fmt_str(xl(ws, 20, 3))
    if raw:
        products = '、'.join(item.split(':')[0] for item in raw.split(';') if item.strip())
    else:
        products = ''
    set_cell(sc(2, 1), products, align='left')

    #r3：行业类别 + 近3个月成交额
    set_cell(sc(3, 1), fv(43, 3))
    set_cell(sc(3, 5), fv(43, 8))

    #r4：标的证券名称 + 近一年最高价
    set_cell(sc(4, 1), fv(44, 3))
    set_cell(sc(4, 5), fv(44, 8))

    #r5-r10：左侧3列数据(c3,c4,c5) + 右侧1列(c8)
    #            Excel行  左c3     左c4     左c5     右c8
    DATA = [(5, 45), (6, 46), (7, 47), (8, 48), (9, 49), (10, 50)]
    for trow, xrow in DATA:
        set_cell(sc(trow, 1), fv(xrow, 3))
        set_cell(sc(trow, 2), fv(xrow, 4))
        set_cell(sc(trow, 3), fv(xrow, 5))
        set_cell(sc(trow, 5), fv(xrow, 8))

    #数据来源段落：替换表格后紧跟的说明文字
    src_text = fmt_str(xl(ws, 52, 2))
    if src_text:
        for p in doc.paragraphs:
            if '数据来源' in p.text and '统计数据' in p.text:
                for run in p.runs: run.clear()
                run = p.runs[0] if p.runs else p.add_run()
                run.text = src_text
                run.font.name = '宋体'; run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                break

    print('标的证券基本面 (标的券概要)')


#退市指标排查 → 往模板"退市风险类型"表格的"标的证券情况"列填值
def fill_delist(doc, wb_val, wb_fmt):
    if '市场' not in wb_val.sheetnames: return print('没有"市场"sheet')
    ws = wb_val['市场']
    wf = wb_fmt['市场']

    #找表
    tbl = None
    for t in doc.tables:
        for row in t.rows[:1]:
            for c in row.cells:
                if '退市风险类型' in c.text: tbl = t; break
            if tbl: break
        if tbl: break
    if tbl is None: return print('未找到"退市风险类型"表格')

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return '-'
        if isinstance(v, str): return v if v.strip() else '-'
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '#,##0' in nf and '0.00' not in nf:
            return f'{v:,.0f}'
        if '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    #Word表行 → Excel行的映射（都取c3）
    MAP = [
        (1, 23),   # 近20个交易日收盘价最低值
        (2, 24),   # 近20个交易日市值最低值
        (3, 25),   # 近120个交易日成交股份数量
        (4, 26),   # 近20个交易日股东数
        (5, 30),   # 近一年审计报告意见类型
        (6, 27),   # 近一年年报收入
        (7, 28),   # 近一年净利润或扣非后净利润孰低
        (8, 29),   # 近一年末净资产
        (9, None),  # 近一年内控审计报告 — Excel没有
        (10, None), # 其他规范类退市风险 — Excel没有
        (11, 31),  # 其他退市风险
    ]
    for trow, xrow in MAP:
        val = fv(xrow, 3) if xrow else '-'
        set_cell(tbl.cell(trow, 2), val)

    print('退市指标排查 (市场)')


#压力测试表 → "二、融资方及一致行动人资产负债情况和压力测试"下方插入
def fill_stress_test(doc, wb_val, wb_fmt):
    if 'Sheet1' not in wb_val.sheetnames: return print('没有"Sheet1"sheet')
    ws = wb_val['Sheet1']
    wf = wb_fmt['Sheet1']

    #找标题段落（跳过目录），再往后找到"……"占位符，在它后面插表
    anchor = None
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or '').lower(): continue
        if '资产负债情况和压力测试' in p.text and not any(k in p.text for k in ['撰写','列示']):
            anchor = p._element; break
    if anchor is None: return print('未找到"资产负债情况和压力测试"段落')

    #跳到"……"占位符后
    el = anchor.getnext()
    insert_after = anchor
    while el is not None:
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(n.text or '' for n in el.iter() if n.tag.endswith('}t')).strip()
            if '……' in txt: insert_after = el; break
            if txt.startswith('（撰写') or txt.startswith('三、'): break
            insert_after = el
        el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return ''
        if isinstance(v, str): return v
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            return f'{v*100:.2f}%'
        if '#,##0' in nf:
            return f'{v:,.2f}'
        return f'{v:,.2f}' if isinstance(v, float) else str(v)

    # 5行4列：表头r1 + 数据r2-r5
    tbl = make_table(doc, 5, 4, insert_after, col_widths=[3.0, 3.0, 3.5, 3.0])
    #表头
    set_cell(tbl.cell(0, 0), '', bold=True)
    for ci, h in enumerate([fv(1,2), fv(1,3), fv(1,4)]):
        set_cell(tbl.cell(0, ci+1), h, bold=True)
    #数据行 r2-r5
    for i, r in enumerate(range(2, 6)):
        set_cell(tbl.cell(i+1, 0), fv(r, 1), align='left')
        set_cell(tbl.cell(i+1, 1), fv(r, 2))
        set_cell(tbl.cell(i+1, 2), fv(r, 3))
        set_cell(tbl.cell(i+1, 3), fv(r, 4))

    print('压力测试 (Sheet1)')


#上市公司概况注册信息 → 往模板"公司名称/注册地址…"表格填值
def fill_company_info(doc, wb_val):
    if '标的券概要' not in wb_val.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb_val['标的券概要']

    #找第一个"公司名称"开头的8行表（即Table 14, 上市公司概况下的表）
    tbl = None
    for t in doc.tables:
        if len(t.rows) == 8 and t.rows[0].cells[0].text.strip() == '公司名称':
            tbl = t; break
    if tbl is None: return print('未找到上市公司概况注册信息表格')

    #Excel c10=标签, c11=值, r17-r25; Word表按标签名匹配填入
    data = {}
    for r in range(17, 26):
        k = xl(ws, r, 10)
        v = xl(ws, r, 11)
        if k: data[str(k).strip()] = fmt_str(v)

    for row in tbl.rows:
        lbl = row.cells[0].text.strip()
        if lbl in data:
            set_cell(row.cells[1], data[lbl], align='left')

    print('上市公司概况 (标的券概要)')


#偿债能力分析表 → "（二）偿债能力分析"下方插入，从标的券概要r110-r128
def fill_solvency(doc, wb_val, wb_fmt):
    if '标的券概要' not in wb_val.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb_val['标的券概要']
    wf = wb_fmt['标的券概要']

    #找正文里的"（二）偿债能力分析"（跳过目录和限售股部分）
    anchor = None
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or '').lower(): continue
        if '限售' in p.text: continue
        if '偿债能力分析' in p.text and p.text.strip().startswith('（二）'):
            anchor = p._element; break
    if anchor is None: return print('未找到"（二）偿债能力分析"段落')

    #跳过撰写说明段落，找插入位置
    el = anchor.getnext()
    insert_after = anchor
    while el is not None:
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(n.text or '' for n in el.iter() if n.tag.endswith('}t')).strip()
            if txt.startswith('（撰写'): insert_after = el
            elif txt.startswith('（三）') or txt.startswith('三、'): break
            elif txt: insert_after = el
        el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return '-'
        if isinstance(v, str): return v if v.strip() else '-'
        if isinstance(v, datetime): return v.strftime('%Y-%m-%d')
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            return f'{v*100:.2f}%'
        if '#,##0' in nf or '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    #行定义：(excel_row, label, is_section_header)
    ROWS = [
        (111, None, False), (112, None, False), (113, None, False),
        (114, None, False), (115, None, False), (116, None, False),
        (117, None, False), (118, None, False), (119, None, False),
        (120, None, False), (121, None, False),
        (122, None, False), (123, None, False), (124, None, False),
        (125, None, True),   # 偿债能力-调整后 子标题
        (126, None, False), (127, None, False), (128, None, False),
    ]

    #表头从r110取日期列
    hdrs = [fv(110, c) for c in [3,4,5,6,7]]
    ncols = 1 + len(hdrs)  # 标签 + 5列数据
    nrows = 1 + len(ROWS)  # 表头 + 数据行

    tbl = make_table(doc, nrows, ncols, insert_after, col_widths=[4.0,2.0,2.0,2.0,2.0,2.0], compact=True)

    #表头行
    set_cell(tbl.cell(0, 0), '偿债能力', bold=True, align='left')
    for ci, h in enumerate(hdrs):
        set_cell(tbl.cell(0, ci+1), h, bold=True, font_size=8)

    #数据行
    for ri, (xr, _, is_hdr) in enumerate(ROWS):
        trow = ri + 1
        label = fv(xr, 2)
        if is_hdr:
            #子标题行：合并全行
            for ci in range(ncols): shade_cell(tbl.cell(trow, ci))
            tbl.cell(trow, 0).merge(tbl.cell(trow, ncols-1))
            set_cell(tbl.cell(trow, 0), label, bold=True, align='left', font_size=8)
        else:
            set_cell(tbl.cell(trow, 0), label, align='left', font_size=8)
            for ci, xc in enumerate([3,4,5,6,7]):
                set_cell(tbl.cell(trow, ci+1), fv(xr, xc), font_size=8)

    print('偿债能力分析 (标的券概要)')


#资产周转率分析 → "（三）资产周转率分析"下方插入存货指标+应收款项指标两张表
def fill_turnover(doc, wb_val, wb_fmt):
    if '标的券概要' not in wb_val.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb_val['标的券概要']
    wf = wb_fmt['标的券概要']

    #找正文"（三）资产周转率分析"（跳过目录）
    anchor = None
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or '').lower(): continue
        if '资产周转率分析' in p.text and p.text.strip().startswith('（三）'):
            anchor = p._element; break
    if anchor is None: return print('未找到"（三）资产周转率分析"段落')

    #跳过撰写说明，找插入位置
    el = anchor.getnext()
    insert_after = anchor
    while el is not None:
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(n.text or '' for n in el.iter() if n.tag.endswith('}t')).strip()
            if txt.startswith('（撰写'): insert_after = el; break
            if txt.startswith('（四）'): break
        el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return '-'
        if isinstance(v, str): return v if v.strip() else '-'
        if isinstance(v, datetime): return v.strftime('%Y-%m-%d')
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf:
            return f'{v*100:.2f}%'
        if '#,##0' in nf or '0.00' in nf:
            return f'{v:.2f}'
        return str(v)

    last = insert_after
    W6 = [4.0, 2.0, 2.0, 2.0, 2.0, 2.0]

    def add_indicator_tbl(title, hdr_row, data_rows):
        nonlocal last
        #小标题段落
        p = make_para(doc, title, last, font_size=10, bold=True); last = p._element
        #表头取日期
        hdrs = [fv(hdr_row, c) for c in [3,4,5,6,7]]
        nr = 1 + len(data_rows)
        t = make_table(doc, nr, 6, last, col_widths=W6, compact=True)
        set_cell(t.cell(0, 0), title, bold=True, align='left', font_size=8)
        for ci, h in enumerate(hdrs):
            set_cell(t.cell(0, ci+1), h, bold=True, font_size=8)
        for ri, xr in enumerate(data_rows):
            set_cell(t.cell(ri+1, 0), fv(xr, 2), align='left', font_size=8)
            for ci, xc in enumerate([3,4,5,6,7]):
                set_cell(t.cell(ri+1, ci+1), fv(xr, xc), font_size=8)
        last = t._element

    #存货指标：r149标题行，r150-155数据行
    add_indicator_tbl('存货指标', 149, [150,151,152,153,154,155])
    print('存货指标')

    #应收款项指标：r163标题行，r164-176数据行（跳过空行r170）
    add_indicator_tbl('应收款项指标', 163, [164,165,166,167,168,169,171,172,173,174,175,176])
    print('应收款项指标')


#现金流分析表 → "（四）现金流分析"下方插入
def fill_cashflow(doc, wb_val, wb_fmt):
    if '标的券概要' not in wb_val.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb_val['标的券概要']
    wf = wb_fmt['标的券概要']

    anchor = None
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or '').lower(): continue
        if '现金流分析' in p.text and p.text.strip().startswith('（四）'):
            anchor = p._element; break
    if anchor is None: return print('未找到"（四）现金流分析"段落')

    #跳过撰写说明
    el = anchor.getnext()
    insert_after = anchor
    while el is not None:
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(n.text or '' for n in el.iter() if n.tag.endswith('}t')).strip()
            if txt.startswith('（撰写'): insert_after = el; break
            if txt.startswith('（五）'): break
        el = el.getnext()

    def fv(r, c):
        v = xl(ws, r, c)
        if v is None: return '-'
        if isinstance(v, str): return v
        if isinstance(v, datetime): return v.strftime('%Y-%m-%d')
        if not isinstance(v, (int, float)): return str(v)
        nf = (wf.cell(r, c).number_format or '')
        if '%' in nf: return f'{v*100:.2f}%'
        if '#,##0' in nf or '0.00' in nf: return f'{v:.2f}'
        return str(v)

    #4行5列：表头 + 经营现金流/净利润/净现比
    hdrs = [fv(104, c) for c in [4,5,6,7]]
    DATA = [(95, fmt2), (99, fmt2), (106, fmt_pct)]

    tbl = make_table(doc, 1+len(DATA), 5, insert_after, col_widths=[4.5,2.5,2.5,2.5,2.5])
    set_cell(tbl.cell(0, 0), '指标名称', bold=True, align='left')
    for ci, h in enumerate(hdrs):
        set_cell(tbl.cell(0, ci+1), h, bold=True)
    for ri, (xr, f) in enumerate(DATA):
        set_cell(tbl.cell(ri+1, 0), fmt_str(xl(ws, xr, 2)), align='left')
        for ci, xc in enumerate([4,5,6,7]):
            set_cell(tbl.cell(ri+1, ci+1), f(xl(ws, xr, xc)))

    print('现金流分析 (标的券概要)')


#  公司对比的基础表格
def fill_comparison(doc, wb):
    if '对标' not in wb.sheetnames: return print('没有"对标"sheet')
    ws = wb['对标']
    anchor = clear_between(doc, '1、对标同行业上市公司的关键指标比较', '2、标的证券核心竞争力')
    if anchor is None: return print('未找到对标段落')

    last = anchor
    W6, W10, W4 = [1.5,2.0,2.5,2.5,2.5,2.5], [1.2,1.3,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5], [1.5,1.5,5.0,5.0]

    def add_tbl(title, headers, data, fsz=9, w=None, compact=False):
        nonlocal last
        p = make_para(doc, title, last, font_size=10, bold=True); last = p._element
        nc, nr = len(headers[-1]), len(headers)+len(data)
        t = make_table(doc, nr, nc, last, col_widths=w, compact=compact)
        for hi, hr in enumerate(headers):
            for ci, h in enumerate(hr): set_cell(t.cell(hi,ci), h, bold=True, font_size=fsz)
        for ri, dr in enumerate(data):
            for ci, v in enumerate(dr):
                set_cell(t.cell(len(headers)+ri,ci), v, font_size=fsz, align='left' if ci<=1 else 'center')
        last = t._element

#生成简单的表格
    def simple6(title, rows, ylabels):
        add_tbl(title, [['代码','证券简称']+ylabels],
            [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in range(5,9)] for r in rows], w=W6)
        print(f'{title}')

#稍微复杂点的表格（两个数据区）
    def dual10(title, hdr, comps):
        yv = [fmt_str(xl(ws,hdr,c)) for c in [5,6,7,8]]
        yg = [fmt_str(xl(ws,hdr,c)) for c in [9,10,11,12]]
        add_tbl(title,
            [['','',f'{title}（亿元）','','','','同比增长率（%）','','',''], ['代码','证券简称']+yv+yg],
            [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in range(5,9)]+[fmt2(xl(ws,r,c)) for c in range(9,13)] for r in comps],
            fsz=8, w=W10, compact=True)
        print(f'{title}')

    #调用顺序：(1)收入构成 → (2)经营规模对比 → (3)盈利能力对比
    #(1) 收入构成
    p = make_para(doc, '（1）收入构成', last, font_size=10, bold=True); last = p._element
    add_tbl('主营构成', [['代码','证券简称','主营构成(产品）','主要产品名称']],
        [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4)),fmt_str(xl(ws,r,5)),fmt_str(xl(ws,r,9))] for r in [25,26,27]], fsz=8, w=W4)
    print('主营构成')

    #(2) 经营规模对比
    p = make_para(doc, '（2）经营规模对比', last, font_size=10, bold=True); last = p._element
    for t,h,c in [('营业总收入',61,[62,63,64]),('净利润',70,[71,72,73]),('归母净利润',81,[82,83,84]),('扣非后归母净利润',90,[91,92,93])]:
        dual10(t,h,c)

    #(3) 盈利能力对比
    p = make_para(doc, '（3）盈利能力对比', last, font_size=10, bold=True); last = p._element
    add_tbl('毛利率',
        [['','','销售毛利率(%)','','','','扣除销售费用后的毛利率(%)','','',''],
         ['代码','证券简称','2023年','2024年','2025年','2026年Q1','2023年','2024年','2025年','2026年Q1']],
        [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in range(5,13)] for r in [103,104,105]],
        fsz=8, w=W10, compact=True)
    print('毛利率')
    simple6('销售净利率(%)',                [41,42,43], ['2023年','2024年','2025年','2026年Q1'])
    simple6('净资产收益率（ROE）(%) -加权', [35,36,37], ['2023年','2024年','2025年','2026年Q1(年化)'])

    #(4) 偿债能力对比：资产负债率+速动比率（10列，结构同毛利率）
    p = make_para(doc, '（4）偿债能力对比', last, font_size=10, bold=True); last = p._element
    add_tbl('资产负债率和速动比率',
        [['','','资产负债率（%）','','','','速动比率（倍）','','',''],
         ['代码','证券简称','2023年','2024年','2025年','2026年Q1','2023年','2024年','2025年','2026年Q1']],
        [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in range(5,13)] for r in [183,184,185]],
        fsz=8, w=W10, compact=True)
    print('资产负债率和速动比率')

    #(5) 经营效率对比：应收账款周转率 + 存货周转率（8列，含最近一期规模）
    W8 = [1.2,1.3,1.5,1.5,1.5,1.5,1.5,1.5]
    p = make_para(doc, '（5）经营效率对比', last, font_size=10, bold=True); last = p._element
    add_tbl('应收账款周转率',
        [['','','应收款项周转率（次/年）','','','','最近一期 应收账款规模',''],
         ['代码','证券简称','2023年','2024年','2025年','2026年Q1(年化)','金额（亿元）','占总资产比例']],
        [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in [5,6,7,8]]+[fmt2(xl(ws,r,9)),fmt_pct(xl(ws,r,11))] for r in [194,195,196]],
        fsz=8, w=W8, compact=True)
    print('应收账款周转率')
    add_tbl('存货周转率',
        [['','','存货周转率（次/年）','','','','最近一期 存货规模',''],
         ['代码','证券简称','2023年','2024年','2025年','2026年Q1(年化)','金额（亿元）','占总资产比例']],
        [[fmt_str(xl(ws,r,3)),fmt_str(xl(ws,r,4))]+[fmt2(xl(ws,r,c)) for c in [5,6,7,8]]+[fmt2(xl(ws,r,9)),fmt_pct(xl(ws,r,11))] for r in [201,202,203]],
        fsz=8, w=W8, compact=True)
    print('存货周转率')


#找关键词删表格，但是要防止误删段落内容
def fill_industry(doc, wb):
    if '标的券概要' not in wb.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb['标的券概要']

    anchor = None
    for p in doc.paragraphs:
        if '与行业上市公司比较' in p.text or '与同行业上市公司比较' in p.text:
            anchor = p._element; break
    if anchor is None: anchor = find_para(doc, '标的券估值风险分析')
    if anchor is None: return print('未找到行业比较段落')

    el = anchor.getnext()
    while el is not None:
        if el.tag.split('}')[-1] == 'tbl': doc.element.body.remove(el); break
        txt = ''.join(n.text or '' for n in el.iter() if n.tag.endswith('}t')).strip()
        if txt and not any(k in txt for k in ['标的券','下表','与']): break
        el = el.getnext()

    #行业比较的数据处理（11个公司）
    comps = []
    for r in range(61, 72):
        if xl(ws,r,3): comps.append([xl(ws,r,c) for c in [2,3,4,6,7,8,9,10]])

    #中位值和平均值：PE_TTM/PB/PB扣商誉直接从Excel读（45-47），PE 26E/27E自己算
    tm  = xl(ws,45,5) or 0;  ta  = xl(ws,45,4) or 0   #PE TTM
    pm  = xl(ws,46,5) or 0;  pa  = xl(ws,46,4) or 0   #PB
    xm  = xl(ws,47,5) or 0;  xa  = xl(ws,47,4) or 0   #PB扣商誉

    valid = [c for c in comps if isinstance(c[3],(int,float)) and 0 < c[3] <= 100]
    def stats(vals):
        f = [x for x in vals if isinstance(x,(int,float))]
        return (statistics.median(f),statistics.mean(f)) if f else (0,0)
    e6m,e6a = stats([c[4] for c in valid if c[4] and c[4]>0])
    e7m,e7a = stats([c[5] for c in valid if c[5] and c[5]>0])

    #格式
    tbl = make_table(doc, 2+len(comps)+2, 8, anchor)
    for ci,h in enumerate(['排名','代码','简称','市盈率PE','市盈率PE','市盈率PE','市净率PB（LF）','市净率（扣除商誉）']):
        set_cell(tbl.cell(0,ci),h,bold=True)
    for ci,h in enumerate(['排名','代码','简称','TTM','26E','27E','市净率PB（LF）','市净率（扣除商誉）']):
        set_cell(tbl.cell(1,ci),h,bold=True)
    for ri,c in enumerate(comps):
        for ci,v in enumerate([fmt0(c[0]),fmt_str(c[1]),fmt_str(c[2]),fmt2(c[3]),fmt2(c[4]),fmt2(c[5]),fmt2(c[6]),fmt2(c[7])]):
            set_cell(tbl.cell(ri+2,ci),v,align='left' if ci in [1,2] else 'center')
    for label,row_i,vals in [('行业中值',2+len(comps),[tm,e6m,e7m,pm,xm]),('行业均值',3+len(comps),[ta,e6a,e7a,pa,xa])]:
        set_cell(tbl.cell(row_i,0),label,bold=True)
        for ci,v in enumerate(vals): set_cell(tbl.cell(row_i,3+ci),fmt2(v))
    print('行业估值比较 (标的券概要)')

def fill_financial(doc, wb):
    if '标的券概要' not in wb.sheetnames: return print('没有"标的券概要"sheet')
    ws = wb['标的券概要']

    # 跳过目录，定位到最后一个匹配的"三...财务分析"
    anchor = None
    for p in doc.paragraphs:
        if 'toc' in (p.style.name or '').lower(): continue
        if '财务分析' in p.text and p.text.strip().startswith('三'):
            anchor = p._element
    if anchor is None: return print('未找到"三、财务分析"段落')

    el = anchor.getnext()
    while el is not None:
        if el.tag.split('}')[-1] == 'tbl': doc.element.body.remove(el); break
        el = el.getnext()

    # Excel行, 标签, 格式化
    R = [(81,'营业收入(亿元)',fmt2),(82,'同比增长率(%)',fmt_pct),(83,'归母净利润(亿元)',fmt2),
         (84,'同比增长率(%)',fmt_pct),(85,'扣非归母净利润(亿元)',fmt2),(86,'同比增长率(%)',fmt_pct),
         (87,'毛利率(%)',fmt_pct),(88,'净资产收益率ROE(加权,%)',fmt_pct),
         (89,'商誉(亿元)',fmt2),(90,'商誉占总资产比例(%)',fmt_pct),
         (91,'总资产(亿元)',fmt2),(92,'归母净资产(亿元)',fmt2),
         (93,'资产负债率(%)',fmt_pct),(94,'速动比率',fmt2),
         (95,'经营活动现金流量净额(亿元)',fmt2),
         (96,'应收账周转率（年化）',fmt2),(97,'存货周转率（年化）',fmt2),
         (99,'净利润(亿元)',fmt2),(100,'同比增长率(%)',fmt_pct),(101,'销售净利润率(%)',fmt_pct)]

    hdrs = [fmt_date(xl(ws,80,c)) if isinstance(xl(ws,80,c),datetime) else str(xl(ws,80,c)) for c in [4,5,6,7]]
    tbl = make_table(doc, 1+len(R), 5, anchor) #建表，填表标题
    set_cell(tbl.cell(0,0),'主要财务科目及比率',bold=True,align='left')
    for ci,h in enumerate(hdrs): set_cell(tbl.cell(0,ci+1),h,bold=True)
    for ri,(xr,label,f) in enumerate(R): #然后填数据
        set_cell(tbl.cell(ri+1,0),label,align='left')
        for ci,xc in enumerate([4,5,6,7]): set_cell(tbl.cell(ri+1,ci+1),f(xl(ws,xr,xc)))
    print('财务分析 (标的券概要)')


#往word模板自带的表里填数据（看表标题关键词判断）
def fill_legacy(doc, wb_val, wb_fmt):
    def find_tbl(kw):
        for t in doc.tables:
            for row in t.rows[:2]:
                for cell in row.cells:
                    if kw in cell.text: return t
        return None
    #读值，再根据根式改成相对应的数字（一次拿数字，一次拿格式）
    def sv(sh, addr):
        try:
            v = wb_val[sh][addr].value; nf = wb_fmt[sh][addr].number_format
            if isinstance(v,datetime): return v.strftime('%Y-%m-%d')
            if isinstance(v,(int,float)): return f'{v:.2%}' if '%' in nf else (f'{v:.2f}' if isinstance(v,float) else str(v))
            return str(v) if v else ''
        except: return None
    #防止崩溃
    def sc(t,r,c):
        try: return t.cell(r,c) if 0<=r<len(t.rows) and 0<=c<len(t.columns) else None
        except: return None

    if '存续' in wb_val.sheetnames:
        t = find_tbl('序号')
        if t:
            while len(t.rows)<9: t.add_row() #防止行数不够
            for cl,wc in {'C':1,'E':3,'F':4,'G':5,'H':6}.items():
                for i in range(1,11): set_cell(sc(t,i,wc), sv('存续',f'{cl}{i+9}'))
            print('存续交易')
    #还没实装下面的东西，但是逻辑有一点了
    if '本次要素' in wb_val.sheetnames:
        t = find_tbl('持股')
        if t:
            d = {}
            for r in range(63,79):
                k = wb_val['本次要素'][f'B{r}'].value
                if k:
                    key = str(k).strip().replace(' ','')
                    d[key] = {'C':sv('本次要素',f'C{r}'), 'E':sv('本次要素',f'E{r}')}
            for row in t.rows:
                lbl = row.cells[0].text.strip().replace(' ','')
                if lbl in d:
                    if len(row.cells)>1: set_cell(row.cells[1], d[lbl].get('C'))
                    if len(row.cells)>4: set_cell(row.cells[4], d[lbl].get('E'))
            print('持股概要')


#  主函数，标题，打开文件拿数据格式，按顺序加载每个模块，保存，以及报错处理
def main():
    print('='*50+'\n  尽调报告自动填表工具\n'+'='*50)
    print(f'工作目录: {BASE_DIR}\n')
    excel_path, word_path = get_files()
    if not excel_path: print('未找到 .xlsx'); input('回车退出...'); return
    if not word_path:   print('未找到 .docx'); input('回车退出...'); return
    print(f'Excel: {os.path.basename(excel_path)}\nWord:  {os.path.basename(word_path)}\n')
    try:
        wb_val = load_workbook(excel_path, data_only=True)
        wb_fmt = load_workbook(excel_path, data_only=False)
        doc = Document(word_path)
        print('【本次要素】'); fill_deal_summary(doc, wb_val, wb_fmt)
        print('\n【持股概要】'); fill_shareholding(doc, wb_val, wb_fmt)
        print('\n【减持受限】'); fill_restrict(doc, wb_val, wb_fmt)
        print('\n【基本面】');   fill_basic_info(doc, wb_val, wb_fmt)
        print('\n【退市排查】'); fill_delist(doc, wb_val, wb_fmt)
        print('\n【压力测试】'); fill_stress_test(doc, wb_val, wb_fmt)
        print('\n【公司概况】'); fill_company_info(doc, wb_val)
        print('\n【偿债能力】'); fill_solvency(doc, wb_val, wb_fmt)
        print('\n【周转率】');   fill_turnover(doc, wb_val, wb_fmt)
        print('\n【现金流】');   fill_cashflow(doc, wb_val, wb_fmt)
        print('\n【对标模块】');   fill_comparison(doc, wb_val)
        print('\n【行业比较】'); fill_industry(doc, wb_val)
        print('\n【财务分析】'); fill_financial(doc, wb_val)
        print('\n【基础模块】'); fill_legacy(doc, wb_val, wb_fmt)
        ts = datetime.now().strftime('%Y%m%d_%H%M')
        out = os.path.join(BASE_DIR, f'合并_尽调报告{ts}.docx')
        doc.save(out)
        print(f'\n保存为: {os.path.basename(out)}')
    except Exception as e:
        print(f'\n出错: {e}'); print(traceback.format_exc())
    input('\n回车关闭...')

if __name__ == '__main__':
    main()
